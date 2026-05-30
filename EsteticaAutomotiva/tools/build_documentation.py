from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Documentacao_Estetica_Automotiva_ABNT.docx"


def set_run(run, bold=False, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run(run)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text.upper() if level == 1 else text)
    set_run(run, bold=True, size=12)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run)
    return p


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for cell, header in zip(tbl.rows[0].cells, headers):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run(r, bold=True, size=10)
    for row in rows:
        cells = tbl.add_row().cells
        for cell, value in zip(cells, row):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run(r, size=10)
    return tbl


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.left_margin = Cm(3)
section.right_margin = Cm(2)
section.bottom_margin = Cm(2)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)
styles["List Bullet"].font.name = "Times New Roman"
styles["List Bullet"].font.size = Pt(12)

for text in [
    "NOME DA INSTITUICAO",
    "CURSO DE ANALISE E DESENVOLVIMENTO DE SISTEMAS",
    "",
    "",
    "SISTEMA WEB PARA ESTETICA AUTOMOTIVA",
    "Controle de clientes e agendamento de servicos com MongoDB Atlas",
    "",
    "",
    "Nome do aluno",
    "",
    "",
    "Cidade",
    "2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run(run, bold=text in {"SISTEMA WEB PARA ESTETICA AUTOMOTIVA"}, size=12)

doc.add_page_break()

heading(doc, "Resumo", 1)
paragraph(
    doc,
    "Este documento apresenta a modelagem e a arquitetura de uma aplicacao Java Web desenvolvida com Maven, Spring Boot e MongoDB Atlas para uma estetica automotiva. O sistema permite cadastro autonomo de clientes, autenticacao, gestao administrativa de clientes, configuracao de tipos de servico e controle de agenda com validacao de disponibilidade.",
)

heading(doc, "1 Introducao", 1)
paragraph(
    doc,
    "A aplicacao foi projetada para organizar o atendimento de uma estetica automotiva, reduzindo conflitos de agenda e permitindo que clientes solicitem servicos sem depender de atendimento manual. A solucao tambem considera principios da Lei Geral de Protecao de Dados, principalmente minimizacao de exposicao, consentimento e controle de acesso.",
)

heading(doc, "2 Requisitos atendidos", 1)
for item in [
    "Login completo com autenticacao por e-mail e senha, controle de perfis e criptografia BCrypt.",
    "Gestao de clientes com cadastro autonomo, edicao, exclusao definitiva e listagem administrativa com dados sensiveis mascarados.",
    "Cadastro e listagem de tipos de servico com nome, descricao detalhada, preco vigente e status ativo ou inativo.",
    "Controle de agenda com visualizacao diaria ou semanal pelo gestor, criacao pelo cliente, edicao e cancelamento.",
    "Validacao para impedir sobreposicao de agendamentos no mesmo horario.",
]:
    bullet(doc, item)

heading(doc, "3 Arquitetura do sistema", 1)
paragraph(
    doc,
    "A arquitetura segue o padrao MVC, separando responsabilidades entre camada de apresentacao, controladores, servicos de negocio, repositorios e documentos MongoDB. As telas foram implementadas com Thymeleaf, enquanto a seguranca e o controle de acesso foram centralizados no Spring Security.",
)
table(
    doc,
    ["Camada", "Responsabilidade"],
    [
        ["View", "Templates Thymeleaf para login, cadastro, clientes, servicos e agenda."],
        ["Controller", "Recebe requisicoes HTTP, valida entradas principais e encaminha fluxos."],
        ["Service", "Aplica regras de negocio, LGPD, mascaramento e conflito de agenda."],
        ["Repository", "Persiste e consulta documentos no MongoDB via Spring Data."],
        ["Model", "Representa usuarios, tipos de servico e agendamentos."],
    ],
)

heading(doc, "4 Modelo de dados", 1)
paragraph(
    doc,
    "O banco MongoDB utiliza colecoes para usuarios, tipos de servico e agendamentos. A entidade AppUser representa tanto clientes quanto gestores, diferenciados pelo atributo role. Appointment registra a relacao entre cliente, tipo de servico, data, hora e status.",
)
table(
    doc,
    ["Classe", "Principais atributos", "Finalidade"],
    [
        ["AppUser", "nome, CPF, telefone, e-mail, senha, placas, perfil, consentimento LGPD", "Armazena clientes e gestor."],
        ["ServiceType", "nome, descricao, preco vigente, ativo", "Define os servicos oferecidos."],
        ["Appointment", "cliente, servico, data, hora, status", "Registra agendamentos."],
    ],
)

heading(doc, "5 Diagrama de classes", 1)
paragraph(doc, "O diagrama de classes contempla os pacotes model, repository e service, evidenciando que um cliente pode possuir varios agendamentos e que um tipo de servico tambem pode estar relacionado a varios agendamentos.")
add_code_block(
    doc,
    [
        "AppUser 1 ---- 0..* Appointment : realiza",
        "ServiceType 1 ---- 0..* Appointment : define",
        "ClientService --> AppUserRepository",
        "AppointmentService --> AppointmentRepository",
        "AppointmentService --> ServiceTypeRepository",
    ],
)

heading(doc, "6 Diagrama de sequencia", 1)
paragraph(doc, "O fluxo principal de agendamento inicia no cliente, passa pela View, chega ao AppointmentController e segue para o AppointmentService. Antes de gravar, o servico consulta os repositorios para verificar o tipo de servico e a disponibilidade do horario.")
add_code_block(
    doc,
    [
        "Cliente -> View: informa servico, data e hora",
        "View -> AppointmentController: POST /agendamentos",
        "AppointmentController -> AppointmentService: schedule(...)",
        "AppointmentService -> Repositories: consulta servico e agenda",
        "AppointmentService -> AppointmentService: valida horario futuro e conflito",
        "AppointmentService -> MongoDB Atlas: salva agendamento",
        "AppointmentController -> View: redireciona para meus agendamentos",
    ],
)

heading(doc, "7 LGPD e seguranca", 1)
for item in [
    "O cadastro exige aceite explicito para tratamento dos dados.",
    "As senhas sao armazenadas com hash BCrypt, nao em texto puro.",
    "A listagem administrativa mascara CPF, e-mail e telefone.",
    "Clientes acessam apenas seus proprios agendamentos; gestor acessa cadastros, servicos e agenda geral.",
    "A connection string do MongoDB Atlas deve ser mantida em variavel de ambiente, fora do codigo-fonte.",
]:
    bullet(doc, item)

heading(doc, "8 Instrucoes de execucao", 1)
paragraph(doc, "Para executar no IntelliJ IDEA, abra a pasta do projeto, aguarde o Maven carregar as dependencias, configure a variavel MONGODB_URI com a connection string do MongoDB Atlas e execute a classe EsteticaAutomotivaApplication.")
add_code_block(doc, ["mvn clean package", "mvn spring-boot:run", "http://localhost:8080"])
paragraph(doc, "O usuario gestor inicial criado automaticamente e admin@estetica.com com senha admin123. Em ambiente real, esta senha deve ser alterada apos o primeiro acesso.")

heading(doc, "9 Consideracoes finais", 1)
paragraph(doc, "O projeto atende aos requisitos principais propostos e pode ser evoluido com recuperacao de senha, auditoria de alteracoes, confirmacao por e-mail, escolha de duracao por tipo de servico e painel de indicadores operacionais.")

doc.save(OUT)
print(OUT)
